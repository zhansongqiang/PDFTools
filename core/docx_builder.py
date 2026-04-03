"""Word 文档生成模块 - 将解析后的内容生成 .docx 文件
Copyright (c) 2026 zsq. All rights reserved.
版本: 1.0
联系邮箱: 11016795@qq.com
联系电话: 18820818283
"""
from __future__ import annotations

import io
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from .pdf_parser import PageContent, TextBlock, ImageBlock, TableBlock


def build_docx(
    pages: list[PageContent],
    output_path: str | Path,
    progress_callback=None,
) -> str:
    """将页面内容生成 Word 文档"""
    doc = Document()
    total_pages = len(pages)

    # 设置默认样式
    _setup_default_style(doc)

    # 第一遍：收集所有页面元素，跨页合并
    all_elements = []
    for page_idx, page in enumerate(pages):
        if progress_callback:
            progress_callback(page_idx + 1, total_pages)

        # 收集所有元素并按 y 坐标排序
        elements = []

        for tb in page.text_blocks:
            elements.append(("text", tb.y, tb))
        for ib in page.image_blocks:
            elements.append(("image", ib.y, ib))
        for tbl in page.table_blocks:
            elements.append(("table", tbl.y, tbl))

        elements.sort(key=lambda e: e[1])
        all_elements.append((page_idx, elements))

    # 第二遍：生成文档
    for page_idx, elements in all_elements:
        for elem_type, _, elem in elements:
            if elem_type == "text":
                _add_text_block(doc, elem)
            elif elem_type == "image":
                _add_image_block(doc, elem)
            elif elem_type == "table":
                _add_table_block(doc, elem)

        # 页面之间添加分页符（最后一页除外）
        if page_idx < total_pages - 1:
            doc.add_page_break()

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return str(output_path)


def _setup_default_style(doc: Document):
    """设置文档默认样式"""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(12)
    # 设置中文字体
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 段落格式
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5  # 1.5倍行距


def _set_run_font(run, font_name: str = "宋体", east_asia: str = None):
    """设置 run 的字体，同时设置东亚字体"""
    run.font.name = font_name
    if east_asia:
        run.element.rPr.rFonts.set(qn('w:eastAsia'), east_asia)
    else:
        run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def _add_text_block(doc: Document, block: TextBlock):
    """添加文本块，保留格式"""
    para = doc.add_paragraph()

    # 设置段落格式
    pf = para.paragraph_format
    pf.line_spacing = 1.5

    if block.is_title:
        _add_title_block(doc, para, block)
    elif block.is_toc_entry:
        _add_toc_block(doc, para, block)
    else:
        _add_body_block(doc, para, block)


def _add_title_block(doc: Document, para, block: TextBlock):
    """添加标题"""
    pf = para.paragraph_format

    if block.title_level == 1 and block.font_size >= 20:
        # 封面/大标题 - 居中，大字号
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_before = Pt(12)
        pf.space_after = Pt(12)
        run = para.add_run(block.text)
        run.font.size = Pt(min(block.font_size, 28))
        run.font.bold = True
        _set_run_font(run, "黑体")
    elif block.title_level == 1:
        # 一级标题（如 "前言", "一、..."）
        pf.space_before = Pt(18)
        pf.space_after = Pt(6)
        run = para.add_run(block.text)
        run.font.size = Pt(min(block.font_size, 18))
        run.font.bold = True
        _set_run_font(run, "黑体")
        if block.alignment == "center":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif block.title_level == 2:
        # 二级标题（如 "（一）..."）
        pf.space_before = Pt(12)
        pf.space_after = Pt(6)
        run = para.add_run(block.text)
        run.font.size = Pt(min(block.font_size, 16))
        run.font.bold = True
        _set_run_font(run, "黑体")
    elif block.title_level == 3:
        # 三级标题（如 "1.语文"）
        pf.space_before = Pt(6)
        pf.space_after = Pt(3)
        run = para.add_run(block.text)
        run.font.size = Pt(min(block.font_size, 14))
        run.font.bold = True
        _set_run_font(run, "黑体")


def _add_toc_block(doc: Document, para, block: TextBlock):
    """添加目录条目 — 保留标题和页码"""
    pf = para.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.15
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 解析目录条目：标题\t页码
    text = block.text
    parts = text.split('\t')
    if len(parts) >= 2:
        title_part = parts[0].strip()
        page_part = parts[-1].strip()
        # 添加标题文字
        run = para.add_run(title_part)
        run.font.size = Pt(10.5)
        _set_run_font(run, "宋体")
        # 添加点线引导符（用制表符+右对齐制表位实现）
        run = para.add_run("\t")
        run.font.size = Pt(10.5)
        # 添加页码
        run = para.add_run(page_part)
        run.font.size = Pt(10.5)
        _set_run_font(run, "宋体")
    else:
        # 没有点线的目录项
        run = para.add_run(text)
        run.font.size = Pt(10.5)
        _set_run_font(run, "宋体")

    # 添加右对齐制表位（带点线引导符）
    _add_tab_stop(para, "right", "dot")


def _add_tab_stop(para, position: str = "right", leader: str = "dot"):
    """为段落添加制表位"""
    from docx.oxml import OxmlElement
    pPr = para._element.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    if position == "right":
        tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '9000')  # 右边界位置
    if leader == "dot":
        tab.set(qn('w:leader'), 'dot')
    tabs.append(tab)
    pPr.append(tabs)


def _add_body_block(doc: Document, para, block: TextBlock):
    """添加正文段落"""
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5

    if block.first_line_indent:
        pf.first_line_indent = Cm(0.74)  # 约2个中文字符

    # 图表标题（小字号，居中）
    if block.font_size <= 11 and block.alignment == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(block.text)
        run.font.size = Pt(block.font_size) if block.font_size > 0 else Pt(10.5)
        _set_run_font(run, "宋体")
        return

    # 普通正文 - 保留混合格式
    if block.spans_info:
        for span_info in block.spans_info:
            run = para.add_run(span_info["text"])
            size = span_info["size"]
            if size > 0:
                run.font.size = Pt(size)
            if span_info["bold"]:
                run.font.bold = True
            if span_info["italic"]:
                run.font.italic = True
            _set_run_font(run, "宋体")
    else:
        run = para.add_run(block.text)
        if block.font_size > 0:
            run.font.size = Pt(block.font_size)
        if block.bold:
            run.font.bold = True
        if block.italic:
            run.font.italic = True
        _set_run_font(run, "宋体")


def _add_image_block(doc: Document, block: ImageBlock):
    """添加图片块"""
    if block.ocr_text:
        # OCR 识别的文字作为文本段落
        for line in block.ocr_text.split("\n"):
            if line.strip():
                para = doc.add_paragraph()
                pf = para.paragraph_format
                pf.first_line_indent = Cm(0.74)
                run = para.add_run(line)
                run.font.size = Pt(12)
                _set_run_font(run, "宋体")
    else:
        # 直接插入图片
        try:
            image_stream = io.BytesIO(block.data)
            width = None
            if block.width > 0:
                # 限制最大宽度为页面宽度的80%
                max_width = 6.0  # Inches
                calc_width = block.width / 72
                width = Inches(min(calc_width, max_width))
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(image_stream, width=width)
        except Exception:
            doc.add_paragraph("[图片无法显示]")


def _add_table_block(doc: Document, block: TableBlock):
    """添加表格，带完整格式"""
    if not block.rows:
        return

    num_cols = max(len(row) for row in block.rows)
    if num_cols == 0:
        return

    table = doc.add_table(rows=len(block.rows), cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置表格宽度为页面宽度
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tblW)

    # 填充数据
    for row_idx, row in enumerate(block.rows):
        for col_idx, cell in enumerate(row):
            if col_idx < num_cols:
                table_cell = table.cell(row_idx, col_idx)
                # 清除默认段落
                for p in table_cell.paragraphs:
                    p.clear()

                # 添加文本
                if table_cell.paragraphs:
                    para = table_cell.paragraphs[0]
                else:
                    para = table_cell.add_paragraph()

                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(cell.text)
                run.font.size = Pt(10.5)
                _set_run_font(run, "宋体")

                # 表头加粗
                if row_idx == 0 or cell.bold:
                    run.font.bold = True

                # 设置单元格垂直居中
                tc = table_cell._tc
                tcPr = tc.get_or_add_tcPr()
                vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
                tcPr.append(vAlign)
